import sqlite3
import sys
import docx

from PyQt5.Qt import QMainWindow, QDialog, QApplication
from PyQt5 import uic
from PyQt5.QtCore import QSize, Qt
from PyQt5.QtWidgets import QTableWidgetItem, QDialog, QDialogButtonBox
from PyQt5.QtGui import QKeyEvent

from docx import Document

SELECT_ALL_PLAYERS = "SELECT name, score FROM players"
SELECT_ALL_TEAMS = "SELECT name, score FROM teams"

DELETE_PLAYER = "DELETE FROM players WHERE name = "

class LoginWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('login.ui', self)
        self.pushButton.clicked.connect(lambda: self.show_main_window())

    def show_main_window(self):
        mainWindow.show()
        self.close()


class Main(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('main_window.ui', self)
        self.show_players_button.clicked.connect(lambda: playersListWindow.show())
        self.show_teams_button.clicked.connect(lambda: teamsListWindow.show())
        self.gen_report_button.clicked.connect(lambda: self.gen_report())

    def gen_report(self):
        doc = Document("report_template.docx")
        doc.tables

        data = sqlite3.connect('teams.db')
        command = data.cursor()
        teams = command.execute("SELECT name, score FROM teams;").fetchall()
        for i in range(len(teams)):
            doc.tables[0].add_row()
            doc.tables[0].cell(i + 1, 0).text = str(i + 1)
            doc.tables[0].cell(i + 1, 1).text = teams[i][0]
            doc.tables[0].cell(i + 1, 2).text = str(teams[i][1])

        data = sqlite3.connect('players.db')
        command = data.cursor()
        players = command.execute("SELECT name, score FROM players;").fetchall()
        for i in range(len(players)):
            doc.tables[1].add_row()
            doc.tables[1].cell(i + 1, 0).text = str(i + 1)
            doc.tables[1].cell(i + 1, 1).text = str(players[i][1])
            doc.tables[1].cell(i + 1, 2).text = players[i][0]

        doc.save("report.docx")


class PlayersListWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('players.ui', self)
        self.data = sqlite3.connect('players.db')
        self.command = self.data.cursor()
        self.result = list()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(['Участник', 'Очки'])
        self.table.setRowCount(0)
        self.result = self.command.execute(SELECT_ALL_PLAYERS).fetchall()
        for i, row in enumerate(self.result):
            self.table.setRowCount(self.table.rowCount() + 1)
            for j, elem in enumerate(row):
                self.table.setItem(i, j, self.createTableItem(str(elem), Qt.ItemIsEnabled | Qt.ItemIsEditable))

        self.add_button.clicked.connect(lambda: addPlayerDialog.show())
        self.delete_button.clicked.connect(lambda: deletePlayerDialog.show())
        self.duplicate_button.clicked.connect(lambda: duplicatePlayerDialog.show())

        self.save_button.clicked.connect(lambda: self.save())

        self.table.cellDoubleClicked.connect(self.onDoubleClick)
        self.table.keyPressEvent = self.onKeyPress

    def createTableItem(self, content, flags):
        item = QTableWidgetItem(content)
        item.setFlags(flags)
        return item

    def onDoubleClick(self, row, col):
        if col == 0:
            item = self.table.item(row, col)
            item.setFlags(Qt.ItemIsEditable)
            item.setFlags(Qt.ItemIsEnabled)

    def onKeyPress(self, key: QKeyEvent):
        if key.key() == Qt.Key_Return:
            self.save()

    def save(self):
        newTable = []
        oldTable = self.result
        for i in range(len(oldTable)):
            name = self.table.item(i, 0).text()
            score = self.table.item(i, 1).text()
            newTable.append((name, score))
        for i, row in enumerate(newTable):
            if row[1] != oldTable[i][1]:
                self.command.execute('UPDATE players SET score = ' + row[1] + ' WHERE name = "' + row[0] + '"')
                self.data.commit()


class TeamsListWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi('teams.ui', self)
        self.data = sqlite3.connect('teams.db')
        self.command = self.data.cursor()
        self.result = list()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(['Название', 'Очки'])
        self.table.setRowCount(0)
        self.result = self.command.execute(SELECT_ALL_TEAMS).fetchall()
        for i, row in enumerate(self.result):
            self.table.setRowCount(self.table.rowCount() + 1)
            for j, elem in enumerate(row):
                self.table.setItem(i, j, self.createTableItem(str(elem), Qt.ItemIsEnabled | Qt.ItemIsEditable))
        self.table.cellDoubleClicked.connect(self.onDoubleClick)

    def createTableItem(self, content, flags):
        item = QTableWidgetItem(content)
        item.setFlags(flags)
        return item

    def onDoubleClick(self, row, col):
        item = self.table.item(row, col)
        item.setFlags(Qt.ItemIsEditable)
        item.setFlags(Qt.ItemIsEnabled)


class AddPlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('add_player.ui', self)


class DeletePlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('delete_player.ui', self)
        self.buttonBox = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel)
        self.buttonBox.accepted.connect(self.accept)
        self.buttonBox.rejected.connect(self.reject)

    def accept(self):
        rows = playersListWindow.table.selectionModel().selectedIndexes()
        for i in range(len(rows)):
            playersListWindow.command.execute(DELETE_PLAYER + '"' + str(playersListWindow.result[rows[i].row()][0]) + '"')
            playersListWindow.table.removeRow(rows[i].row())
        playersListWindow.data.commit()
        self.close()


class DuplicatePlayerDialog(QDialog):
    def __init__(self):
        super().__init__()
        uic.loadUi('duplication.ui', self)


app = QApplication(sys.argv)
loginWindow = LoginWindow()
loginWindow.show()

mainWindow = Main()

playersListWindow = PlayersListWindow()
teamsListWindow = TeamsListWindow()

addPlayerDialog = AddPlayerDialog()
deletePlayerDialog = DeletePlayerDialog()
duplicatePlayerDialog = DuplicatePlayerDialog()
sys.exit(app.exec_())
